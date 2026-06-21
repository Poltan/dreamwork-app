"""
Dreamwork backend — FastAPI.

Endpoints:
  POST /api/match   multipart: resume file + prefs -> ranked jobs + parsed profile
  POST /api/letter  json: {job, profile, resume_text, lang} -> tailored cover letter
  GET  /            serves the frontend (../frontend/index.html)
  GET  /api/health  liveness + which providers/keys are configured

Run:
  pip install -r requirements.txt
  cp .env.example .env   # fill in keys
  uvicorn app:app --reload --port 8000
"""

import os
import io
import json
import time
import threading
import pathlib
from collections import defaultdict, deque
from concurrent.futures import ThreadPoolExecutor

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, Response

try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    pass

import providers

ANTHROPIC_MODEL = os.getenv("ANTHROPIC_MODEL", "claude-sonnet-4-6")
# Fast model for the structured resume-parsing step (big speed win, negligible quality risk).
FAST_MODEL = os.getenv("FAST_MODEL", "claude-haiku-4-5-20251001")
# Model for job ranking. Ranking is a structured pick-top-N task, so the fast model
# handles it well and roughly halves the AI time on the critical path. Override with
# RANK_MODEL=claude-sonnet-4-6 to go back to the heavier model if quality needs it.
RANK_MODEL = os.getenv("RANK_MODEL", FAST_MODEL)

# --- YooKassa payments (paid cover letters) ----------------------------------
# Whole paywall is gated behind PAID_LETTERS so the code can ship safely (letters
# stay free) and be switched on only after testing with YooKassa test keys.
PAID_LETTERS = os.getenv("PAID_LETTERS") == "1"
LETTER_PRICE = os.getenv("LETTER_PRICE", "99.00")
# Paid product #2 — resume generation (improve an uploaded resume or build from a form).
PAID_RESUME = os.getenv("PAID_RESUME") == "1"
RESUME_PRICE = os.getenv("RESUME_PRICE", "1999.00")
YK_SHOP_ID = os.getenv("YOOKASSA_SHOP_ID")
YK_SECRET = os.getenv("YOOKASSA_SECRET_KEY")
YK_RETURN_URL = os.getenv("YOOKASSA_RETURN_URL", "https://www.dreamworkjob.ru/?paid=1")
# Send receipt (НПД чек) data with the payment — YooKassa forms the self-employed cheque.
YK_RECEIPT = os.getenv("YOOKASSA_RECEIPT", "1") == "1"
# One payment unlocks exactly one letter. In-memory is enough for a single instance.
_PAID_LOCK = threading.Lock()
_USED_PAYMENTS = set()

_HERE = pathlib.Path(__file__).resolve().parent
# Works for both layouts: flat (index.html next to app.py, used on the host/Render)
# and structured (../frontend/index.html, used locally).
FRONTEND = _HERE / "index.html"
if not FRONTEND.exists():
    FRONTEND = _HERE.parent / "frontend" / "index.html"

app = FastAPI(title="Dreamwork API", version="0.1")
# The frontend is served from this same service, so we don't need a wildcard CORS.
# Restrict to our own origin (+ localhost for dev) so other websites can't drive our
# paid endpoints from a user's browser. Override with ALLOWED_ORIGINS (comma-separated).
_ORIGINS = [o.strip() for o in os.getenv(
    "ALLOWED_ORIGINS",
    "https://www.dreamworkjob.ru,https://dreamworkjob.ru,https://dreamwork-0nmr.onrender.com,http://localhost:8000,http://127.0.0.1:8000",
).split(",") if o.strip()]
app.add_middleware(
    CORSMiddleware, allow_origins=_ORIGINS, allow_methods=["*"], allow_headers=["*"],
)

# Basic security headers on every response.
@app.middleware("http")
async def _security_headers(request: Request, call_next):
    resp = await call_next(request)
    resp.headers["X-Content-Type-Options"] = "nosniff"
    resp.headers["X-Frame-Options"] = "DENY"
    resp.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
    resp.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
    return resp

# Expose internal debug fields only when explicitly enabled.
DEBUG_API = os.getenv("DEBUG_API") == "1"

# ---- Lightweight in-memory rate limiting (per client IP) -------------------
# Protects the paid endpoints (Anthropic / provider quotas) from abuse. Single
# instance => in-memory is enough; resets on restart.
_RL_LOCK = threading.Lock()
_RL_HITS = defaultdict(deque)

def _client_ip(request: Request) -> str:
    xff = request.headers.get("x-forwarded-for", "")
    if xff:
        return xff.split(",")[0].strip()
    return request.client.host if request.client else "?"

def _rate_limit(bucket: str, ip: str, limit: int, window: int):
    """Allow `limit` hits per `window` seconds per (bucket, ip). Raise 429 if exceeded."""
    now = time.time()
    key = f"{bucket}:{ip}"
    with _RL_LOCK:
        dq = _RL_HITS[key]
        while dq and dq[0] <= now - window:
            dq.popleft()
        if len(dq) >= limit:
            retry = int(window - (now - dq[0])) + 1
            raise HTTPException(429, detail={"code": "rate_limited", "retry": retry})
        dq.append(now)
        # opportunistic cleanup to bound memory
        if len(_RL_HITS) > 5000:
            for k in [k for k, v in list(_RL_HITS.items()) if not v]:
                _RL_HITS.pop(k, None)

# Max resume upload size (bytes). Bounds memory + blocks upload-bomb DoS.
MAX_UPLOAD_BYTES = int(os.getenv("MAX_UPLOAD_BYTES", str(5 * 1024 * 1024)))


def _anthropic():
    key = os.getenv("ANTHROPIC_API_KEY")
    if not key:
        raise HTTPException(500, "ANTHROPIC_API_KEY is not set on the server.")
    import anthropic
    return anthropic.Anthropic(api_key=key)


def _ask_claude(prompt, max_tokens=1500, system=None, model=None):
    client = _anthropic()
    msg = client.messages.create(
        model=model or ANTHROPIC_MODEL,
        max_tokens=max_tokens,
        system=system or "You are a precise assistant. Follow instructions exactly.",
        messages=[{"role": "user", "content": prompt}],
    )
    return "".join(b.text for b in msg.content if getattr(b, "type", None) == "text").strip()


def _extract_json(text):
    text = text.strip()
    if text.startswith("```"):
        text = text.strip("`")
        text = text[text.find("{") if "{" in text else text.find("["):]
    for op, cl in (("{", "}"), ("[", "]")):
        i, j = text.find(op), text.rfind(cl)
        if i != -1 and j != -1 and j > i:
            try:
                return json.loads(text[i:j + 1])
            except Exception:
                pass
    return json.loads(text)


# Caps to keep parsing of untrusted files bounded (anti zip-bomb / huge-doc DoS).
MAX_PDF_PAGES = 30
MAX_DOCX_PARAS = 4000
MAX_TEXT_CHARS = 60000

# Only these file types are accepted as resumes. We validate BOTH the extension
# and the file's binary signature (magic bytes) so a renamed executable / binary
# can't slip through just by being called "resume.pdf".
ALLOWED_EXTS = (".pdf", ".docx", ".txt")
# Signatures of common non-text/executable formats we explicitly reject for .txt.
_BINARY_HEADS = (b"MZ", b"\x7fELF", b"\xca\xfe\xba\xbe", b"\xfe\xed\xfa", b"PK", b"%PDF")


def _err(status: int, code: str, **extra):
    """Raise an HTTPException whose body carries a machine-readable `code`,
    so the frontend can show a friendly, localized message (never raw text)."""
    detail = {"code": code}
    detail.update(extra)
    raise HTTPException(status, detail=detail)


def parse_resume(filename: str, data: bytes) -> str:
    name = (filename or "").lower().strip()
    ext = ("." + name.rsplit(".", 1)[-1]) if "." in name else ""
    if ext not in ALLOWED_EXTS:
        _err(415, "bad_format")

    head = data[:8]

    if ext == ".pdf":
        if not head.startswith(b"%PDF"):
            _err(415, "bad_format")
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                pages = pdf.pages[:MAX_PDF_PAGES]
                text = "\n".join((p.extract_text() or "") for p in pages)
        except Exception:
            _err(400, "pdf_unreadable")
        return text[:MAX_TEXT_CHARS]

    if ext == ".docx":
        # A real .docx is a ZIP container, which always starts with the "PK" signature.
        if not head.startswith(b"PK"):
            _err(415, "bad_format")
        try:
            import docx
            d = docx.Document(io.BytesIO(data))
            text = "\n".join(p.text for p in d.paragraphs[:MAX_DOCX_PARAS])
        except Exception:
            _err(400, "docx_unreadable")
        return text[:MAX_TEXT_CHARS]

    # .txt — reject binaries (null bytes or executable/container signatures), then decode.
    if b"\x00" in data[:4096] or head.startswith(_BINARY_HEADS):
        _err(415, "bad_format")
    for enc in ("utf-8", "cp1251", "latin-1"):
        try:
            return data.decode(enc)[:MAX_TEXT_CHARS]
        except Exception:
            continue
    _err(415, "bad_format")


# Rough FX rates to RUB for cross-currency salary comparison.
_FX = {"RUB": 1.0, "USD": 90.0, "EUR": 95.0, "GBP": 110.0, "ILS": 25.0, "SAR": 24.0,
       "AED": 24.0, "CAD": 66.0, "AUD": 60.0, "PLN": 23.0, "INR": 1.1, "BRL": 18.0,
       "MXN": 5.0, "NZD": 55.0, "SGD": 67.0, "ZAR": 5.0}

def _to_rub(amount, currency):
    try:
        amount = float(amount)
    except Exception:
        return 0.0
    return amount * _FX.get((currency or "RUB").upper(), 1.0)


def _profile_is_meaningful(p) -> bool:
    """True only if we extracted enough from the resume to drive a real job search.
    Blocks the 'random text still returns jobs' case (empty goals -> generic search)."""
    if not isinstance(p, dict):
        return False
    if p.get("is_resume") is False:
        return False
    has_titles = any((t or "").strip() for t in (p.get("target_titles") or []))
    has_query = bool((p.get("search_query") or "").strip()
                     or (p.get("search_query_en") or "").strip())
    has_skills = any((s or "").strip() for s in (p.get("skills") or []))
    return has_titles or has_query or has_skills


PROFILE_PROMPT = """Извлеки из резюме структурированный профиль кандидата.
ВАЖНО: сначала определи, действительно ли это резюме/CV (а не рассказ, статья, случайный текст,
список покупок и т.п.). Если перед тобой НЕ резюме и из текста нельзя понять профессию и цель
поиска работы — верни {{"is_resume": false}} и оставь остальные поля пустыми (НЕ придумывай профессию).
Верни ТОЛЬКО JSON без пояснений, по схеме:
{{
  "is_resume": true,
  "name": "",
  "headline": "краткая роль одной строкой",
  "country": "страна кандидата по локации/языку резюме, на английском (напр. Russia, USA, Israel); если не ясно — пусто",
  "seniority": "junior|mid|senior|lead|director|c-level",
  "target_titles": ["3-6 реалистичных названий вакансий; если у кандидата есть редкий язык или сильная региональная/отраслевая экспертиза (напр. китайский язык, опыт с Китаем), добавь 1-2 варианта, отражающих это (напр. 'Project Manager China', 'ВЭД-менеджер Китай')"],
  "search_query": "1-2 СЛОВА — основное название роли для поиска вакансий (на языке рынка кандидата), без навыков и названий программ",
  "search_query_en": "1-3 English keywords (role translated to English) for international search, e.g. 'project manager'",
  "skills": ["8-15 ключевых навыков"],
  "industries": ["отрасли кандидата"],
  "languages": ["языки кандидата"],
  "desired_salary": "желаемая зарплата из резюме — ТОЛЬКО число без пробелов/валюты, или null",
  "salary_currency": "валюта желаемой зарплаты (RUB/USD/EUR/CNY/...), или пусто",
  "summary": "2-3 предложения о кандидате, включая ключевые преимущества (языки, отраслевая/региональная экспертиза)"
}}

РЕЗЮМЕ:
{resume}"""


RANK_PROMPT = """Ты — карьерный консультант. Ниже профиль кандидата и список вакансий.
Отбери НЕ БОЛЕЕ {top} самых подходящих вакансий: по роли, грейду, отрасли, локации и зарплатным ожиданиям ({salary_note}).
Если страна поиска не задана — допускай сильные вакансии из разных стран мира (география может быть разной).
Сильные стороны кандидата (редкие языки — напр. китайский; отраслевая или региональная экспертиза — напр. опыт работы с Китаем) — весомый плюс: вакансии, где они прямо востребованы, поднимай выше.
Лучше вернуть 2-3 точных совпадения, чем добивать список слабыми.
Для каждой выбранной верни оценку соответствия и короткую заметку (1-2 предложения, на языке профиля), почему вакансия подходит именно этому кандидату.
Не выдумывай вакансии, используй только id из списка. Отсей нерелевантные (другая профессия/грейд).

Верни ТОЛЬКО JSON-массив:
[{{"id":"<id вакансии>","score":0-100,"fit":"<заметка>"}}]

ПРОФИЛЬ:
{profile}

ВАКАНСИИ:
{jobs}"""


LETTER_PROMPT = """Ты — профессиональный консультант по сопроводительным письмам с 15+ годами опыта.
Напиши короткое, цепляющее, «человеческое» сопроводительное письмо под конкретную вакансию.
Сначала про себя проанализируй вакансию (5-10 требований, домен, ключевые слова) — анализ НЕ выводи.

Правила: {lang_rule} 3-5 абзацев по 1-3 предложения, ~700-1200 знаков; живой стиль без канцелярита и штампов
(«обладаю навыками», «успешно реализовывал», «дружная команда», «данная вакансия» — запрещены); 8-15 ключевых слов
из вакансии естественно; зеркаль 2-3 формулировки работодателя; не дублируй резюме; без выдуманных цифр.
Первый абзац — позиция + хук. Середина — 2-3 конкретных акцента опыта под задачи вакансии. Финал — спокойный
call-to-action одним предложением. Выведи ТОЛЬКО текст письма с приветствием и подписью кандидата.

КАНДИДАТ (профиль): {profile}

РЕЗЮМЕ (фрагмент): {resume}

ВАКАНСИЯ: {title} — {company} ({location}).
ОПИСАНИЕ: {desc}"""


def _valid_email(e: str) -> bool:
    e = (e or "").strip()
    return bool(e) and 3 <= len(e) <= 254 and "@" in e and "." in e.split("@")[-1]


def _yk_configured() -> bool:
    return bool(YK_SHOP_ID and YK_SECRET)


def _yk_request(method: str, path: str, body=None):
    """Call the YooKassa API v3 with HTTP Basic auth (shopId:secretKey)."""
    import requests as _rq
    import base64 as _b64
    import uuid as _uuid
    auth = _b64.b64encode(f"{YK_SHOP_ID}:{YK_SECRET}".encode()).decode()
    headers = {"Authorization": f"Basic {auth}", "Content-Type": "application/json"}
    if method == "POST":
        headers["Idempotence-Key"] = str(_uuid.uuid4())
    return _rq.request(method, f"https://api.yookassa.ru/v3{path}",
                       json=body, headers=headers, timeout=20)


def _payment_ok(payment_id: str, product: str) -> bool:
    """A payment is valid for `product` only if it succeeded AND was created for that
    product (so a 99₽ letter payment can't unlock a 1999₽ resume, and vice versa)."""
    try:
        r = _yk_request("GET", f"/payments/{payment_id}")
        if r.status_code != 200:
            return False
        d = r.json()
        if not (d.get("status") == "succeeded" and bool(d.get("paid"))):
            return False
        meta_product = (d.get("metadata") or {}).get("product")
        if product == "letter":
            # accept legacy/missing tags for backward compatibility
            return meta_product in (None, "", "letter", "cover_letter")
        return meta_product == product
    except Exception:
        return False


@app.get("/api/health")
def health():
    return {
        "ok": True,
        "model": ANTHROPIC_MODEL,
        "paid_letters": PAID_LETTERS,
        "letter_price": LETTER_PRICE,
        "paid_resume": PAID_RESUME,
        "resume_price": RESUME_PRICE,
        "pay_configured": _yk_configured(),
        "keys": {
            "anthropic": bool(os.getenv("ANTHROPIC_API_KEY")),
            "jooble": bool(os.getenv("JOOBLE_API_KEY")),
            "adzuna": bool(os.getenv("ADZUNA_APP_ID") and os.getenv("ADZUNA_APP_KEY")),
            "superjob": bool(os.getenv("SUPERJOB_KEY")),
            "hh_token": bool(os.getenv("HH_TOKEN")),
            "scraper": bool(os.getenv("SCRAPER_API_KEY")),
            "jsearch": bool(os.getenv("RAPIDAPI_KEY")),
            "company_ats": bool(os.getenv("COMPANY_ATS")),
            "proxy": bool(os.getenv("PROXY_URL")),
        },
        "always_on": ["Habr Career", "Remotive", "Trudvsem"],
    }


@app.get("/api/proxytest")
def proxytest(key: str = ""):
    # Diagnostic endpoint — gated behind DIAG_KEY so it doesn't publicly leak the
    # egress/proxy IP. Set DIAG_KEY in the environment and call /api/proxytest?key=...
    diag = os.getenv("DIAG_KEY")
    if not diag or key != diag:
        raise HTTPException(404, "Not found")
    import requests as _rq
    u = os.getenv("PROXY_URL")
    px = {"http": u, "https": u} if u else None
    out = {"proxy_set": bool(u)}
    try:
        r = _rq.get("https://api.ipify.org?format=json", proxies=px, timeout=25)
        out["egress_ip"] = r.json().get("ip")
    except Exception as e:
        out["ip_error"] = str(e)[:200]
    try:
        r2 = _rq.get("https://api.hh.ru/vacancies", params={"text": "менеджер", "per_page": 1},
                     proxies=px, timeout=30,
                     headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"})
        out["hh_status"] = r2.status_code
        out["hh_body"] = r2.text[:150]
    except Exception as e:
        out["hh_error"] = str(e)[:200]
    return out


@app.post("/api/match")
async def match(
    request: Request,
    resume: UploadFile = File(...),
    country: str = Form(""),
    location: str = Form(""),
    salary_min: str = Form(""),
    currency: str = Form(""),
    remote_ok: str = Form("false"),
    top: int = Form(10),
):
    _rate_limit("match", _client_ip(request), limit=15, window=600)
    top = max(1, min(int(top), 12))

    data = await resume.read()
    if not data:
        _err(400, "empty_file")
    if len(data) > MAX_UPLOAD_BYTES:
        _err(413, "too_large", max_mb=MAX_UPLOAD_BYTES // (1024 * 1024))
    resume_text = parse_resume(resume.filename, data)
    if len(resume_text.strip()) < 50:
        _err(400, "no_text")

    try:
        profile = _extract_json(_ask_claude(
            PROFILE_PROMPT.format(resume=resume_text[:7000]), max_tokens=1500, model=FAST_MODEL))
    except Exception:
        profile = {}

    # Guard against "garbage in, garbage out": if the upload isn't a real resume
    # (a story, random text, an image-only scan), the model can't extract a profile.
    # Don't silently fall back to a generic search — tell the user clearly.
    if not _profile_is_meaningful(profile):
        _err(422, "no_profile")

    # If the user left the country blank, infer it from the resume (location/language).
    if not (country or "").strip():
        country = (profile.get("country") or "").strip()

    sal = int(salary_min) if str(salary_min).strip().isdigit() else None
    eff_currency = (currency or "RUB").upper()
    # If the user didn't set a salary, use the one stated in the resume.
    if sal is None and profile.get("desired_salary"):
        try:
            sal = int(float(profile.get("desired_salary")))
            eff_currency = (profile.get("salary_currency") or "RUB").upper()
        except Exception:
            sal = None

    # Build query variants. For non-RU markets prefer English keywords/titles.
    code_now = providers.normalize_country(country)
    intl = bool(code_now) and code_now != "RU"
    titles = [(t or "").strip() for t in (profile.get("target_titles") or []) if (t or "").strip()]
    if intl:
        titles.sort(key=lambda t: 0 if all(ord(c) < 128 for c in t) else 1)  # English titles first
    sq = (profile.get("search_query") or "").strip()
    sq_en = (profile.get("search_query_en") or "").strip()
    primary = sq_en if (intl and sq_en) else sq
    candidates = []
    if primary:
        candidates.append(primary)
    for t in titles[:4]:
        if t not in candidates:
            candidates.append(t)
    if sq and sq not in candidates:
        candidates.append(sq)
    hl = (profile.get("headline") or "").strip()
    if hl and hl not in candidates:
        candidates.append(hl)
    # Last-resort query from skills (profile is already validated as meaningful above).
    if not candidates:
        skills = [s for s in (profile.get("skills") or []) if (s or "").strip()]
        if skills:
            candidates.append(skills[0].strip())
    # If we still have nothing to search with, treat it as an undetectable profile
    # rather than returning unrelated jobs.
    if not candidates:
        _err(422, "no_profile")

    # Run a couple of query variants IN PARALLEL; aggregate + de-dupe.
    # Capped at 2 to bound memory (each variant fans out to all providers).
    queries = candidates[:2]
    remote_flag = str(remote_ok).lower() == "true"
    jobs, debug, seen_ids = [], {}, set()

    def _q(q):
        return providers.search_jobs(keywords=q, location=location, country=country,
                                     salary_min=sal, remote_ok=remote_flag, per_provider=15)

    with ThreadPoolExecutor(max_workers=max(1, len(queries))) as ex:
        for qjobs, qdebug in ex.map(_q, queries):
            debug.update(qdebug)
            for j in qjobs:
                if j["id"] in seen_ids:
                    continue
                seen_ids.add(j["id"])
                jobs.append(j)

    # Unified salary filter: convert to RUB, keep matches AND jobs with unstated salary.
    if sal:
        min_rub = _to_rub(sal, eff_currency)
        kept = []
        for j in jobs:
            top_sal = j.get("salary_max") or j.get("salary_min")
            if not top_sal or _to_rub(top_sal, j.get("currency")) >= min_rub * 0.9:
                kept.append(j)
        # Soft filter: apply only if enough remain; otherwise keep the full pool so a
        # high salary ask doesn't zero out results (salary stays a ranking preference).
        if len(kept) >= 3:
            jobs = kept

    if not jobs:
        empty = {"profile": profile, "jobs": [],
                 "note": "Провайдеры не вернули вакансий по этому запросу."}
        if DEBUG_API:
            empty["debug"] = debug; empty["tried"] = candidates
        return JSONResponse(empty)

    # Prioritize company-direct (ATS) jobs. They're fetched last, so without this they
    # sit at the tail of the pool and get crowded out of both the AI ranking window
    # (first 40) and the final list — defeating the "straight from employer" value.
    # Put a bounded slice of them up front so they actually compete and can surface.
    _ats = [j for j in jobs if str(j.get("source", "")).startswith("Company")]
    if _ats:
        _other = [j for j in jobs if not str(j.get("source", "")).startswith("Company")]
        jobs = _ats[:20] + _other + _ats[20:]

    trimmed = [{"id": j["id"], "title": j["title"], "company": j["company"],
                "location": j["location"], "salary": j["salary"],
                "remote": j["remote"], "desc": (j["description"] or "")[:300]}
               for j in jobs[:40]]
    salary_note = f"от {sal} {eff_currency}" if sal else "не указаны"
    try:
        ranked = _extract_json(_ask_claude(
            RANK_PROMPT.format(top=top, salary_note=salary_note,
                               profile=json.dumps(profile, ensure_ascii=False),
                               jobs=json.dumps(trimmed, ensure_ascii=False)),
            max_tokens=1500, model=RANK_MODEL))
    except Exception:
        ranked = [{"id": j["id"], "score": None, "fit": ""} for j in jobs[:top]]

    # The model sometimes returns an object instead of a JSON array — normalize.
    if isinstance(ranked, dict):
        ranked = next((v for v in ranked.values() if isinstance(v, list)), [])
    if not isinstance(ranked, list):
        ranked = []

    by_id = {j["id"]: j for j in jobs}
    results, used = [], set()
    for r in ranked[:top]:
        j = by_id.get(r.get("id"))
        if not j or j["id"] in used:
            continue
        used.add(j["id"])
        j = dict(j)
        j["score"] = r.get("score")
        j["fit"] = r.get("fit", "")
        results.append(j)
    # Append remaining pool jobs (unranked extras) for the "show more" button.
    for j in jobs:
        if len(results) >= 25:
            break
        if j["id"] in used:
            continue
        used.add(j["id"])
        j = dict(j)
        j["score"] = None
        j["fit"] = ""
        results.append(j)

    out = {"profile": profile, "jobs": results,
           "resume_excerpt": resume_text[:4000]}
    if DEBUG_API:
        out["debug"] = debug
        out["pool_by_source"] = {}
        for j in jobs:
            out["pool_by_source"][j.get("source", "?")] = out["pool_by_source"].get(j.get("source", "?"), 0) + 1
    return out


@app.post("/api/letter")
async def letter(payload: dict, request: Request):
    _rate_limit("letter", _client_ip(request), limit=30, window=600)

    # Paywall: when enabled, a letter requires a confirmed, unused YooKassa payment.
    if PAID_LETTERS:
        pid = (payload.get("payment_id") or "").strip()
        if not pid:
            _err(402, "payment_required")
        with _PAID_LOCK:
            if pid in _USED_PAYMENTS:
                _err(409, "payment_used")
        if not _payment_ok(pid, "letter"):
            _err(402, "payment_not_confirmed")
        with _PAID_LOCK:
            if pid in _USED_PAYMENTS:
                _err(409, "payment_used")
            _USED_PAYMENTS.add(pid)

    job = payload.get("job") or {}
    profile = payload.get("profile") or {}
    resume = (payload.get("resume_text") or "")[:4000]
    lang = (payload.get("lang") or "ru").lower()
    lang_rule = ("Пиши письмо на английском, те же правила." if lang == "en"
                 else "Пиши письмо на русском.")
    text = _ask_claude(LETTER_PROMPT.format(
        lang_rule=lang_rule,
        profile=json.dumps(profile, ensure_ascii=False),
        resume=resume,
        title=job.get("title", ""), company=job.get("company", ""),
        location=job.get("location", ""), desc=(job.get("description") or "")[:1500],
    ), max_tokens=1200)
    return {"letter": text}


@app.post("/api/pay")
async def pay(payload: dict, request: Request):
    """Create a YooKassa payment for one cover letter; return a redirect URL."""
    _rate_limit("pay", _client_ip(request), limit=20, window=600)
    if not _yk_configured():
        _err(503, "pay_unconfigured")
    email = (payload.get("email") or "").strip()
    if YK_RECEIPT and not _valid_email(email):
        _err(400, "email_required")
    product = (payload.get("product") or "letter").lower()
    if product not in ("letter", "resume"):
        product = "letter"
    if product == "resume":
        price, desc = RESUME_PRICE, "Резюме под вакансию — DreamWork"
    else:
        price, desc = LETTER_PRICE, "Сопроводительное письмо под вакансию — DreamWork"
    # Return the user to whatever domain they came from (works for both the onrender
    # URL and the custom domain dreamworkjob.ru); fall back to the configured default.
    _host = request.headers.get("host")
    _proto = request.headers.get("x-forwarded-proto", "https")
    return_url = f"{_proto}://{_host}/?paid=1" if _host else YK_RETURN_URL
    body = {
        "amount": {"value": price, "currency": "RUB"},
        "capture": True,
        "confirmation": {"type": "redirect", "return_url": return_url},
        "description": desc,
        "metadata": {"product": product},
    }
    if YK_RECEIPT:
        body["receipt"] = {
            "customer": {"email": email},
            "items": [{
                "description": desc[:128],
                "quantity": "1.00",
                "amount": {"value": price, "currency": "RUB"},
                "vat_code": 1,            # без НДС (самозанятый/НПД)
                "payment_subject": "service",
                "payment_mode": "full_payment",
            }],
        }
    try:
        r = _yk_request("POST", "/payments", body)
    except Exception:
        _err(502, "pay_error")
    if r.status_code not in (200, 201):
        _err(502, "pay_error")
    data = r.json()
    return {
        "payment_id": data.get("id"),
        "confirmation_url": (data.get("confirmation") or {}).get("confirmation_url"),
        "status": data.get("status"),
    }


@app.get("/api/pay/status")
def pay_status(payment_id: str = ""):
    if not _yk_configured():
        _err(503, "pay_unconfigured")
    if not payment_id.strip():
        _err(400, "bad_request")
    try:
        r = _yk_request("GET", f"/payments/{payment_id.strip()}")
    except Exception:
        _err(502, "pay_error")
    if r.status_code != 200:
        _err(502, "pay_error")
    d = r.json()
    return {"status": d.get("status"), "paid": bool(d.get("paid"))}


RESUME_IMPROVE_PROMPT = """Ты — эксперт по составлению резюме с 15+ годами опыта в рекрутинге.
Перепиши и усиль резюме кандидата, сохраняя правдивость (НЕ выдумывай факты, компании, цифры).
Структура: Контакты · Краткое резюме (2-3 предложения) · Опыт работы (по местам, с достижениями и измеримыми результатами, где они есть) · Ключевые навыки · Образование · Языки (если есть).
Стиль: сильные глаголы действия, конкретика, без воды и канцелярита (общие фразы вроде «коммуникабельный», «стрессоустойчивый» убери или замени фактами). Пиши {lang_rule}.
Целевая роль/вакансия: {target} — адаптируй акценты под неё, если она указана.
Верни ТОЛЬКО готовый текст резюме, без пояснений.

ИСХОДНОЕ РЕЗЮМЕ:
{resume}"""

RESUME_CREATE_PROMPT = """Ты — эксперт по составлению резюме. Собери профессиональное резюме из данных кандидата.
Используй только предоставленные факты (НЕ выдумывай опыт, компании, цифры); если данных мало — сделай аккуратное резюме из того, что есть.
Структура: Контакты · Краткое резюме · Опыт работы · Ключевые навыки · Образование · Языки.
Стиль: сильные формулировки, измеримые результаты где уместно, без воды. Пиши {lang_rule}.
Верни ТОЛЬКО готовый текст резюме, без пояснений.

ДАННЫЕ КАНДИДАТА (JSON):
{fields}"""


@app.post("/api/resume")
async def resume(payload: dict, request: Request):
    _rate_limit("resume", _client_ip(request), limit=20, window=600)

    # Paywall: when enabled, a resume requires a confirmed, unused resume payment.
    if PAID_RESUME:
        pid = (payload.get("payment_id") or "").strip()
        if not pid:
            _err(402, "payment_required")
        with _PAID_LOCK:
            if pid in _USED_PAYMENTS:
                _err(409, "payment_used")
        if not _payment_ok(pid, "resume"):
            _err(402, "payment_not_confirmed")
        with _PAID_LOCK:
            if pid in _USED_PAYMENTS:
                _err(409, "payment_used")
            _USED_PAYMENTS.add(pid)

    lang = (payload.get("lang") or "ru").lower()
    lang_rule = "на английском языке" if lang == "en" else "на русском языке"
    mode = (payload.get("mode") or "improve").lower()
    if mode == "create":
        fields = payload.get("fields") or {}
        if not any(str(v).strip() for v in fields.values()):
            _err(400, "no_fields")
        prompt = RESUME_CREATE_PROMPT.format(
            lang_rule=lang_rule, fields=json.dumps(fields, ensure_ascii=False)[:6000])
    else:
        rtext = (payload.get("resume_text") or "")[:7000]
        if len(rtext.strip()) < 30:
            _err(400, "no_text")
        target = (payload.get("target") or "").strip()[:200]
        prompt = RESUME_IMPROVE_PROMPT.format(
            lang_rule=lang_rule, target=(target or "не указана"), resume=rtext)

    text = _ask_claude(prompt, max_tokens=1800)
    return {"resume": text}


import re as _re_md

def _resume_to_docx(text: str) -> bytes:
    """Render the markdown-ish resume text into a .docx (headings, bullets, bold)."""
    import docx
    d = docx.Document()

    def add_runs(p, s):
        for part in _re_md.split(r"(\*\*.+?\*\*)", s):
            if part.startswith("**") and part.endswith("**") and len(part) > 4:
                p.add_run(part[2:-2]).bold = True
            elif part:
                p.add_run(part)

    for raw in (text or "").split("\n"):
        s = raw.strip()
        if not s or _re_md.match(r"^-{3,}$", s):
            continue
        if s.startswith("### "):
            d.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            d.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            d.add_heading(s[2:], level=1)
        elif _re_md.match(r"^[-*•]\s+", s):
            add_runs(d.add_paragraph(style="List Bullet"), _re_md.sub(r"^[-*•]\s+", "", s))
        else:
            add_runs(d.add_paragraph(), s)

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


@app.post("/api/resume/docx")
async def resume_docx(payload: dict, request: Request):
    """Turn a generated resume (markdown text) into a downloadable .docx."""
    _rate_limit("resume_docx", _client_ip(request), limit=30, window=600)
    text = (payload.get("resume") or "").strip()
    if len(text) < 20:
        _err(400, "no_text")
    try:
        data = _resume_to_docx(text)
    except Exception:
        _err(500, "server")
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="resume.docx"'},
    )


@app.post("/api/parse")
async def parse(request: Request, resume: UploadFile = File(...)):
    """Extract text from an uploaded resume (used by the 'improve resume' flow so we
    have the text before redirecting to payment). Same validation as /api/match."""
    _rate_limit("parse", _client_ip(request), limit=20, window=600)
    data = await resume.read()
    if not data:
        _err(400, "empty_file")
    if len(data) > MAX_UPLOAD_BYTES:
        _err(413, "too_large", max_mb=MAX_UPLOAD_BYTES // (1024 * 1024))
    text = parse_resume(resume.filename, data)
    if len(text.strip()) < 50:
        _err(400, "no_text")
    return {"resume_text": text[:7000]}


def _static_file(name: str):
    p = _HERE / name
    if not p.exists():
        p = _HERE.parent / "frontend" / name
    return p


@app.get("/oferta")
def oferta():
    p = _static_file("oferta.html")
    if p.exists():
        return FileResponse(str(p))
    return JSONResponse({"error": "oferta.html not found"}, status_code=404)


@app.get("/privacy")
def privacy():
    p = _static_file("privacy.html")
    if p.exists():
        return FileResponse(str(p))
    return JSONResponse({"error": "privacy.html not found"}, status_code=404)


@app.get("/consent")
def consent():
    p = _static_file("consent.html")
    if p.exists():
        return FileResponse(str(p))
    return JSONResponse({"error": "consent.html not found"}, status_code=404)


@app.get("/")
def index():
    if FRONTEND.exists():
        return FileResponse(str(FRONTEND))
    return JSONResponse({"error": "frontend/index.html not found"}, status_code=404)
